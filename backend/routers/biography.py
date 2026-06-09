"""
人物传记生成 API Router

端点:
  POST /biography/start           - 启动传记生成流程
  POST /biography/step/{step_id}  - 执行单个步骤 (BIO01~BIO05)
  GET  /biography/status/{sid}    - 查询传记生成进度
  GET  /biography/result/{sid}    - 获取最终传记
  GET  /biography/preview/{sid}   - 获取传记摘要预览
"""
from pathlib import Path
import asyncio
import hashlib
import time
from typing import Any, Dict, Optional
from urllib.parse import quote
import threading
from fastapi.responses import Response, StreamingResponse

from fastapi import APIRouter, HTTPException, Query
from pydantic import BaseModel, Field
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup
import io
import re
import base64
from typing import Optional
from copy import deepcopy
# 统一导入业务服务与文件存储
from services import service_manager, session_store
from core import storage as core_storage
from llm_client import call_skill
from skill_loader import load_skill

ROOT_DIR = Path(__file__).resolve().parent.parent.parent
SKILLS_DIR = ROOT_DIR / "skills"

router = APIRouter(prefix="/biography", tags=["biography"])


# ── 请求/响应模型 ─────────────────────────────────────────────────────
class BioStartRequest(BaseModel):
    """启动传记生成请求"""
    sid: Optional[str] = Field(None, description="现有 session ID；若为空则创建新 session")
    form_data: Optional[Dict[str, Any]] = Field(None, description="基础表单数据（deceased_name/birth_date/death_date 等）")


class BioStepRequest(BaseModel):
    """执行单步骤请求"""
    sid: str = Field(..., description="session ID")


class BioStartResponse(BaseModel):
    """启动响应"""
    ok: bool
    session_id: str
    bio_state: Dict[str, Any]
    message: str = ""


class BioStepResponse(BaseModel):
    """步骤执行响应"""
    ok: bool
    step: str
    duration_sec: float
    message: str = ""


class BioStatusResponse(BaseModel):
    """状态查询响应"""
    session_id: str
    step_status: Dict[str, str]
    current_progress: str
    current_step: str
    status: str
    paused: bool = False
    canceled: bool = False
    progress: float
    steps_completed: int
    total_steps: int
    error: Optional[str] = None


class BioResultResponse(BaseModel):
    """最终结果响应"""
    ok: bool
    biography_final: str
    bio_css: str = ""
    biography_json: Dict[str, Any]
    quality_assessment: Dict[str, Any]
    info_gaps: list
    timeline: list

class BioControlResponse(BaseModel):
    """暂停/取消/恢复响应"""
    ok: bool
    session_id: str
    paused: bool
    canceled: bool
    status: str

# ── 端点实现 ─────────────────────────────────────────────────────────

@router.post("/start", response_model=BioStartResponse, summary="启动传记生成流程")
def start_biography(req: BioStartRequest):
    """
    启动人物传记生成流程。
    
    - 若 `sid` 为空，创建新 session
    - 若提供 `form_data`，填充基础表单
    - 返回 session_id 和初始 bio_state
    """
    try:
        sid = req.sid
        form_data = req.form_data or {}
        if not sid:
            sid = session_store.create_session(form_data=form_data)
        else:
            if req.form_data:
                session_store.patch_form(sid, req.form_data)
        if form_data:
            form_data["birth_date"] = _year_only(form_data.get("birth_date", "") or "")
            form_data["death_date"] = _year_only(form_data.get("death_date", "") or "")
        
        s = session_store.require(sid)
        form = s.get("form_data", {}) or {}
        user_id = form.get("user_id")
        memorial_id = form.get("memorial_id")
        if user_id and memorial_id:
            try:
                person = core_storage.ensure_memorial_for_person(
                    user_id,
                    form.get("deceased_name", "") or "未命名",
                    _year_only(form.get("birth_date", "") or ""),
                    _year_only(form.get("death_date", "") or ""),
                    relation=form.get("relation", "") or "",
                    note=form.get("family_memory_text", "")[:200] if form.get("family_memory_text") else "",
                )
                if person:
                    memorial_id = person.get("memorial_id", memorial_id)
                    s["form_data"]["memorial_id"] = memorial_id
                    core_storage.ensure_memorial_assets(user_id, memorial_id)
                    core_storage.update_memorial_meta(user_id, memorial_id, {
                        "name": form.get("deceased_name", "未命名") or "未命名",
                        "birth_date": _year_only(form.get("birth_date", "") or ""),
                        "death_date": _year_only(form.get("death_date", "") or ""),
                        "relation": form.get("relation", "") or "",
                        "note": form.get("family_memory_text", "")[:200] if form.get("family_memory_text") else "",
                    })
            except Exception as ex:
                print(f"[biography] memorial ensure/reuse failed: {ex}")
        if "bio_state" not in s or not s["bio_state"]:
            s["bio_state"] = {
                "current_step": "BIO01",           
                "steps_completed": 0,              
                "total_steps": 6,                   
                "status": "running",                
                "step_status": {                    
                    "BIO01": "pending",
                    "BIO02": "pending",
                    "BIO03": "pending",
                    "BIO04": "pending",
                    "BIO05": "pending",
                    "BIO06": "pending"
                },
                "error": None
            }
        print(f"[biography] started session={sid}")
        return BioStartResponse(
            ok=True,
            session_id=sid,
            bio_state=s["bio_state"],
            message="传记生成流程已启动",
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

def _year_only(value: str) -> str:
    import re
    text = (value or "").strip()
    m = re.search(r"(\d{4})", text)
    return m.group(1) if m else text

@router.post("/pause/{sid}", response_model=BioControlResponse, summary="暂停传记生成")
def pause_biography(sid: str):
    try:
        s = session_store.require(sid)
        bio_state = s["bio_state"]
        control = bio_state.setdefault("control", {"paused": False, "canceled": False})
        control["paused"] = True
        control["canceled"] = False
        return BioControlResponse(
            ok=True,
            session_id=sid,
            paused=True,
            canceled=False,
            status="paused",
        )
    except KeyError:
        raise HTTPException(status_code=404, detail=f"session not found: {sid}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/cancel/{sid}", response_model=BioControlResponse, summary="取消传记生成")
def cancel_biography(sid: str):
    try:
        s = session_store.require(sid)
        bio_state = s["bio_state"]
        control = bio_state.setdefault("control", {"paused": False, "canceled": False})
        control["paused"] = False
        control["canceled"] = True
        current_running = next((step for step, state in bio_state["step_status"].items() if state == "running"), None)
        if current_running:
            bio_state["step_status"][current_running] = "cancelled"
        return BioControlResponse(
            ok=True,
            session_id=sid,
            paused=False,
            canceled=True,
            status="canceled",
        )
    except KeyError:
        raise HTTPException(status_code=404, detail=f"session not found: {sid}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/resume/{sid}", response_model=BioControlResponse, summary="恢复传记生成")
def resume_biography(sid: str):
    try:
        s = session_store.require(sid)
        bio_state = s["bio_state"]
        control = bio_state.setdefault("control", {"paused": False, "canceled": False})
        control["paused"] = False
        control["canceled"] = False
        for step, state in bio_state["step_status"].items():
            if state == "cancelled":
                bio_state["step_status"][step] = "pending"
        return BioControlResponse(
            ok=True,
            session_id=sid,
            paused=False,
            canceled=False,
            status="running",
        )
    except KeyError:
        raise HTTPException(status_code=404, detail=f"session not found: {sid}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/step/{step_id}", response_model=BioStepResponse, summary="执行单个传记生成步骤")
def execute_bio_step(step_id: str, req: BioStepRequest):
    """
    执行单个传记生成步骤（BIO01~BIO05）。
    
    - BIO01: 素材信息提取
    - BIO02: 信息审核与去重
    - BIO03: 时间线重建
    - BIO04: 传记文本生成（核心）
    - BIO05: 质量评审与润色
    
    前置依赖会自动检查。
    """
    try:
        result = service_manager.run_bio_step(req.sid, step_id.upper())
        
        if result.get("error"):
            raise HTTPException(
                status_code=400,
                detail=result.get("message", "步骤执行失败"),
            )
        
        return BioStepResponse(
            ok=True,
            step=step_id.upper(),
            duration_sec=result.get("duration_sec", 0),
            message=f"{step_id} 执行完成",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/chain", response_model=dict, summary="串行执行所有传记生成步骤")
def execute_bio_chain(req: BioStepRequest):
    """
    一键启动完整的传记生成流程（BIO01 → BIO06）。

    返回所有步骤的执行结果、最终传记与排版结果。
    """
    try:
        print(f"[biography] chain requested session={req.sid}")
        result = service_manager.run_bio_chain(req.sid)
        print(f"[biography] chain completed session={req.sid} error={result.get('error', False)}")
        
        if result.get("error"):
            raise HTTPException(
                status_code=400,
                detail=result.get("message", "传记生成流程失败"),
            )
        
        return {
            "ok": True,
            "message": "传记生成完成",
            "biography_final": result.get("biography_final", ""),
            "results": result.get("results", {}),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/status/{sid}", response_model=BioStatusResponse, summary="查询传记生成进度")
def get_bio_status(sid: str):
    """
    查询指定 session 的传记生成进度。
    
    返回各步骤的执行状态（pending/running/approved/error）。
    """
    try:
        s = session_store.require(sid)
        bio_state = s["bio_state"]
        step_status = bio_state["step_status"]
        
        # 计算已完成的步骤数
        steps_completed = sum(1 for v in step_status.values() if v == "approved")
        total_steps = len(step_status)
        
        # 确定当前进度
        control = bio_state.setdefault("control", {"paused": False, "canceled": False})
        if control.get("canceled"):
            current_progress = "已取消"
            status = "canceled"
        elif control.get("paused"):
            current_progress = "已暂停"
            status = "paused"
        elif steps_completed == total_steps and total_steps > 0:
            current_progress = "完成"
            status = "completed"
        elif any(v == "error" for v in step_status.values()):
            current_progress = "失败"
            status = "failed"
        elif any(v == "running" for v in step_status.values()):
            current_progress = "处理中"
            status = "running"
        else:
            current_progress = "等待中"
            status = "pending"

        current_step = next(
            (step for step, state in step_status.items() if state == "running"),
            None,
        )
        if not current_step:
            current_step = next(
                (step for step, state in step_status.items() if state == "pending"),
                "完成",
            )

        progress = steps_completed / total_steps if total_steps else 0.0
        error_msg = bio_state.get("last_error") if any(v == "error" for v in step_status.values()) else None
        progress = steps_completed / total_steps if total_steps else 0.0
        # current_step = next(
        #     (step for step, state in step_status.items() if state == "running"),
        #     next((step for step, state in step_status.items() if state == "pending"), "完成")
        # )
        print(f"[biography] status session={sid} progress={current_progress} completed={steps_completed}/{total_steps} status={step_status}")
        return BioStatusResponse(
            session_id=sid,
            step_status=step_status,
            current_progress=current_progress,
            current_step=current_step,
            status=status,
            paused=control.get("paused", False),
            canceled=control.get("canceled", False),
            progress=progress,
            steps_completed=steps_completed,
            total_steps=total_steps,
            error=error_msg,
        )
    except KeyError:
        raise HTTPException(status_code=404, detail=f"session not found: {sid}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/result/{sid}", response_model=BioResultResponse, summary="获取最终传记")
def get_bio_result(sid: str, embed_images: bool = Query(False, description="是否嵌入图片base64")):
    """
    获取最终传记，可选是否将图片转为base64嵌入（用于PDF导出）
    """
    try:
        result = service_manager.get_biography_result(sid)
        
        if result.get("error"):
            raise HTTPException(
                status_code=400,
                detail=result.get("message") or result.get("detail") or "传记还未生成完成",
            )
        
        biography_final = result.get("biography_final", "")
        
        if embed_images:
            # 方案A：转为 base64 嵌入（推荐用于 PDF）
            biography_final = convert_images_to_base64(biography_final, sid)
        else:
            # 方案B：转为带签名的 URL
            biography_final = convert_images_to_signed_urls(biography_final, sid)
        
        bio_css = result.get("bio_css") or ""
        if not bio_css:
            try:
                s = session_store.require(sid)
                assets = s.get("assets", []) or []
            except Exception:
                assets = []
            bio_css = generate_biography_css(biography_final, assets)
        return BioResultResponse(
            ok=True,
            biography_final=biography_final,
            bio_css=bio_css,
            biography_json=result.get("biography_json", {}),
            quality_assessment=result.get("quality_assessment", {}),
            info_gaps=result.get("info_gaps", []),
            timeline=result.get("timeline", []),
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def convert_images_to_base64(markdown_text: str, sid: str) -> str:
    """将Markdown中的图片转为base64嵌入"""
    import re
    import base64
    
    def replace_with_base64(match):
        alt_text = match.group(1)
        image_path = match.group(2)
        
        # 提取 asset_id
        if '/api/memorials/' in image_path:
            parts = image_path.split('/')
            if 'assets' in parts:
                asset_idx = parts.index('assets')
                if asset_idx + 1 < len(parts):
                    asset_id = parts[asset_idx + 1].split('?')[0]
                    
                    # 从 session 获取 memorial_id
                    s = session_store.require(sid)
                    memorial_id = s.get("form_data", {}).get("memorial_id")
                    
                    if memorial_id and asset_id:
                        # 获取图片文件路径
                        user_id = s.get("form_data", {}).get("user_id")
                        if user_id:
                            asset_path = core_storage.get_asset_path(user_id, memorial_id, asset_id)
                            if asset_path and asset_path.exists():
                                try:
                                    with open(asset_path, 'rb') as f:
                                        image_data = f.read()
                                        base64_str = base64.b64encode(image_data).decode('utf-8')
                                        
                                        # 判断图片类型
                                        if asset_path.suffix.lower() in ['.jpg', '.jpeg']:
                                            mime_type = 'image/jpeg'
                                        elif asset_path.suffix.lower() == '.png':
                                            mime_type = 'image/png'
                                        elif asset_path.suffix.lower() == '.gif':
                                            mime_type = 'image/gif'
                                        elif asset_path.suffix.lower() == '.webp':
                                            mime_type = 'image/webp'
                                        else:
                                            mime_type = 'image/jpeg'
                                        
                                        data_url = f"data:{mime_type};base64,{base64_str}"
                                        return f"![{alt_text}]({data_url})"
                                except Exception as e:
                                    print(f"图片转base64失败: {asset_path}, error={e}")
        
        # 如果转换失败，返回原样
        return match.group(0)
    
    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    return re.sub(pattern, replace_with_base64, markdown_text)

def convert_images_to_signed_urls(markdown_text: str, sid: str) -> str:
    """将Markdown中的图片转为可公开访问的传记图片URL。"""
    import re
    import time

    def replace_with_signed_url(match):
        alt_text = match.group(1)
        image_path = match.group(2)

        if image_path.startswith('/api/memorials/'):
            parts = image_path.split('/')
            if 'assets' not in parts:
                return match.group(0)
            try:
                asset_idx = parts.index('assets')
                memorial_idx = parts.index('memorials')
                memorial_id = parts[memorial_idx + 1]
                asset_id = parts[asset_idx + 1].split('?')[0]
                s = session_store.require(sid)
                user_id = s.get("form_data", {}).get("user_id")
                if not user_id or not memorial_id or not asset_id:
                    return match.group(0)
                timestamp = int(time.time())
                signature = hashlib.md5(f"/api/biography/image/{user_id}/{memorial_id}/{asset_id}{timestamp}secret_key".encode('utf-8')).hexdigest()[:16]
                new_path = f"/api/biography/image/{user_id}/{memorial_id}/{asset_id}?t={timestamp}&s={signature}"
                return f"![{alt_text}]({new_path})"
            except Exception:
                return match.group(0)

        return match.group(0)

    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    return re.sub(pattern, replace_with_signed_url, markdown_text)

@router.get("/preview/{sid}", response_model=dict, summary="获取传记大白话预览")
def get_bio_preview(sid: str):
    """
    获取人物传记的简短摘要预览（大白话）。
    
    用于 UI 展示，约 200~300 字。
    """
    try:
        s = session_store.require(sid)
        bio_state = s["bio_state"]
        
        if bio_state["step_status"].get("BIO04") != "approved":
            return {
                "error": True,
                "message": "传记还未生成（至少需要 BIO04 完成）",
            }
        
        # 从 bio_json 中提取简单预览
        bio_json = bio_state.get("bio_json", {})
        sections = bio_json.get("sections", [])
        
        # 组合前几个章节作为预览
        preview_parts = []
        for section in sections[:3]:  # 引言、早年经历、人生历程
            content = section.get("content", "")[:100]  # 取每个章节的前100字
            if content:
                preview_parts.append(content)
        
        preview = "...".join(preview_parts)[:300]
        
        return {
            "ok": True,
            "preview": preview,
            "word_count": bio_json.get("word_count", 0),
            "emotional_tone": bio_json.get("emotional_tone", ""),
        }
    except KeyError:
        raise HTTPException(status_code=404, detail=f"session not found: {sid}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/image/{user_id}/{memorial_id}/{asset_id}")
def get_biography_image(
    user_id: str, 
    memorial_id: str, 
    asset_id: str,
    t: int = Query(..., description="时间戳"),
    s: str = Query(..., description="签名")
):
    """
    获取传记中的图片（带签名验证，用于PDF导出）
    """
    from fastapi.responses import FileResponse

    expected_path = f"/api/biography/image/{user_id}/{memorial_id}/{asset_id}"
    expected_signature = hashlib.md5(f"{expected_path}{t}secret_key".encode('utf-8')).hexdigest()[:16]

    if s != expected_signature:
        raise HTTPException(status_code=403, detail="Invalid signature")

    if time.time() - t > 86400:
        raise HTTPException(status_code=403, detail="Link expired")

    asset_meta = next((a for a in core_storage.list_assets(user_id, memorial_id) if a.get("asset_id") == asset_id), None)
    if not asset_meta:
        raise HTTPException(status_code=404, detail="Image not found")

    stored_name = asset_meta.get("stored_name") or ""
    asset_path = core_storage.memorial_dir(user_id, memorial_id) / "assets" / stored_name
    if not asset_path.exists():
        raise HTTPException(status_code=404, detail="Image not found")

    return FileResponse(asset_path)
def _build_pdf_html(markdown_text: str, css_text: str = "") -> str:
    body_html = markdown_to_html(markdown_text)
    css_block = f"<style>{css_text}</style>" if css_text else ""
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta http-equiv="Content-Type" content="text/html; charset=UTF-8">
        <style>{generate_biography_css()}</style>
        {css_block}
        <style>
        @page {{
            size: A4;
            margin: 2cm 2cm 2.2cm 2cm;
            @bottom-center {{
                content: "念念传记 · " counter(page);
                font-size: 9pt;
                color: #777;
            }}
        }}
        html, body {{
            -webkit-print-color-adjust: exact;
            print-color-adjust: exact;
        }}
        footer {{
            font-size: 9pt;
            color: #888;
            text-align: center;
            margin-top: 40px;
            border-top: 1px solid #eee;
            padding-top: 10px;
        }}
        </style>
    </head>
    <body class="bio-prose-classic">
        {body_html}
        <footer>
            <p>生成于念念传记写作平台 · 珍藏永恒记忆</p>
        </footer>
    </body>
    </html>
    """


async def _render_pdf_with_playwright(html: str) -> bytes:
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=['--no-sandbox', '--disable-setuid-sandbox']
        )
        try:
            page = await browser.new_page(viewport={"width": 1240, "height": 1754}, device_scale_factor=1)
            await page.set_content(html, wait_until='networkidle')
            pdf_bytes = await page.pdf(
                format='A4',
                print_background=True,
                margin={"top": "20mm", "right": "20mm", "bottom": "22mm", "left": "20mm"}
            )
            return pdf_bytes
        finally:
            await browser.close()


@router.get("/export-pdf/{sid}", summary="导出传记为PDF（Playwright）")
def export_biography_pdf(
    sid: str,
    embed_images: bool = Query(True, description="是否嵌入图片")
):
    """
    使用 Playwright 浏览器打印导出 PDF
    """
    try:
        result = service_manager.get_biography_result(sid)
        if result.get("error"):
            raise HTTPException(status_code=400, detail=result.get("message") or result.get("detail") or "传记未完成")

        biography_md = result.get("biography_final", "")
        if embed_images:
            biography_md = convert_images_to_base64(biography_md, sid)

        html = _build_pdf_html(biography_md)
        pdf_bytes = asyncio.run(_render_pdf_with_playwright(html))

        s = session_store.require(sid)
        deceased_name = s.get("form_data", {}).get("deceased_name", "传记")
        filename = f"{deceased_name}_个人传记.pdf"
        disposition = f"attachment; filename*=UTF-8''{quote(filename)}"
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": disposition},
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"[biography] PDF export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _set_doc_background_white(doc: Document):
    for section in doc.sections:
        sect_pr = section._sectPr
        for child in list(sect_pr.findall(qn('w:background'))):
            sect_pr.remove(child)
    styles = doc.styles
    for style_name in ['Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3']:
        if style_name in styles:
            st = styles[style_name]
            if hasattr(st, 'font'):
                st.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


@router.get("/export-docx/{sid}", summary="导出传记为DOCX（python-docx）")
def export_biography_docx(
    sid: str,
    embed_images: bool = Query(True, description="是否嵌入图片")
):
    """
    使用 python-docx 将传记导出为 Word 文档
    """
    try:
        result = service_manager.get_biography_result(sid)
        if result.get("error"):
            raise HTTPException(status_code=400, detail=result.get("message") or result.get("detail") or "传记未完成")

        biography_md = result.get("biography_final", "")
        if embed_images:
            biography_md = convert_images_to_base64(biography_md, sid)

        doc = create_word_document(biography_md, sid)
        _set_doc_background_white(doc)

        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)

        s = session_store.require(sid)
        deceased_name = s.get("form_data", {}).get("deceased_name", "传记")
        filename = f"{deceased_name}_个人传记.docx"
        disposition = f"attachment; filename*=UTF-8''{quote(filename)}"
        return StreamingResponse(
            io.BytesIO(docx_file.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": disposition},
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"[biography] DOCX export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def markdown_to_html(markdown_text: str) -> str:
    """将 Markdown 转换为 HTML"""
    md = markdown.Markdown(extensions=['extra', 'codehilite', 'toc'])
    html = md.convert(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')

    for table in soup.find_all('table'):
        table['class'] = table.get('class', []) + ['bio-table']

    for img in soup.find_all('img'):
        if img.parent.name != 'div' or 'bio-image' not in img.parent.get('class', []):
            wrapper = soup.new_tag('div', **{'class': 'bio-image'})
            img.wrap(wrapper)
            if img.get('alt'):
                caption = soup.new_tag('div', **{'class': 'bio-image-caption'})
                caption.string = img['alt']
                wrapper.append(caption)

    return str(soup)


def add_pdf_styles(html_content: str) -> str:
    """为 PDF 添加样式"""
    css = """
    <style>
        @page {
            size: A4;
            margin: 2cm 2cm 2.2cm 2cm;
            @bottom-center {
                content: "念念传记 · " counter(page);
                font-size: 9pt;
                color: #777;
            }
        }
        html, body {
            font-family: "Noto Sans CJK SC", "Source Han Sans SC", "WenQuanYi Micro Hei", "Microsoft YaHei", sans-serif;
            line-height: 1.78;
            color: #222;
            font-size: 12pt;
            -webkit-print-color-adjust: exact;
            print-color-adjust: exact;
        }
        body {
            margin: 0;
            padding: 0;
        }
        h1, h2, h3 {
            text-align: center;
            margin-left: auto;
            margin-right: auto;
            page-break-after: avoid;
        }
        h1 { font-size: 24pt; margin-top: 18px; margin-bottom: 14px; color: #222; }
        h2 { font-size: 18pt; margin-top: 24px; margin-bottom: 12px; color: #333; }
        h3 { font-size: 14pt; margin-top: 18px; margin-bottom: 10px; color: #444; }
        p { margin: 0 0 12px 0; text-align: justify; orphans: 2; widows: 2; }
        strong { font-weight: 700; }
        em { font-style: italic; }
        code {
            font-family: "JetBrains Mono", "Consolas", monospace;
            background: #f5f7fa;
            border-radius: 4px;
            padding: 0 4px;
        }
        blockquote {
            margin: 14px 0;
            padding: 10px 16px;
            border-left: 4px solid #d7b27a;
            background: #fbf7ef;
            color: #555;
            font-style: italic;
        }
        .bio-image {
            margin: 18px auto;
            text-align: center;
            page-break-inside: avoid;
        }
        .bio-image img {
            display: block;
            max-width: 100%;
            width: auto;
            height: auto;
            margin: 0 auto;
            border-radius: 8px;
        }
        .bio-image-caption {
            text-align: center;
            color: #666;
            font-size: 9.5pt;
            font-style: italic;
            margin-top: 6px;
        }
        ul, ol { margin: 0 0 12px 20px; }
        li { margin: 0 0 6px 0; }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 16px 0;
            page-break-inside: avoid;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 8px 10px;
            text-align: left;
        }
        th { background: #f6f6f6; }
        hr {
            border: none;
            border-top: 1px solid #ddd;
            margin: 20px 0;
        }
        footer {
            font-size: 9pt;
            color: #888;
            text-align: center;
            margin-top: 40px;
            border-top: 1px solid #eee;
            padding-top: 10px;
        }
    </style>
    """
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta http-equiv="Content-Type" content="text/html; charset=UTF-8">
        <title>个人传记</title>
        {css}
    </head>
    <body>
        {html_content}
        <footer>
            <p>生成于念念传记写作平台 · 珍藏永恒记忆</p>
        </footer>
    </body>
    </html>
    """


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def _clear_paragraph_format(paragraph):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5


def _set_run_font(run, name='微软雅黑', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_inline_html_runs(paragraph, node):
    if isinstance(node, str):
        if node:
            run = paragraph.add_run(node)
            _set_run_font(run)
        return
    name = getattr(node, 'name', None)
    if name is None:
        text = node.get_text() if hasattr(node, 'get_text') else str(node)
        if text:
            run = paragraph.add_run(text)
            _set_run_font(run)
        return
    if name == 'strong':
        run = paragraph.add_run(node.get_text())
        _set_run_font(run, bold=True)
        return
    if name == 'em':
        run = paragraph.add_run(node.get_text())
        _set_run_font(run, italic=True)
        return
    if name == 'code':
        run = paragraph.add_run(node.get_text())
        _set_run_font(run, name='Consolas', size=10)
        return
    if name == 'a':
        run = paragraph.add_run(node.get_text())
        _set_run_font(run, color='0563C1')
        run.underline = True
        return
    if name == 'br':
        paragraph.add_run().add_break()
        return
    for child in getattr(node, 'children', []):
        _add_inline_html_runs(paragraph, child)


def _add_image_to_doc(doc: Document, image_url: str, alt_text: str = '', width_inches: float = 5.8):
    if not (image_url.startswith('data:image') and ',' in image_url):
        return
    try:
        base64_data = image_url.split(',', 1)[1]
        image_data = base64.b64decode(base64_data)
        img_stream = io.BytesIO(image_data)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_paragraph_format(p)
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(width_inches))
        if alt_text:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _clear_paragraph_format(cap)
            cap_run = cap.add_run(alt_text)
            _set_run_font(cap_run, size=9, italic=True, color='666666')
    except Exception as e:
        print(f"图片写入 DOCX 失败: {e}")


def _prepare_docx(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = '微软雅黑'
            st._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def create_word_document(markdown_text: str, sid: str) -> Document:
    """使用 python-docx 创建 Word 文档"""
    doc = Document()
    _prepare_docx(doc)

    html = markdown_to_html(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')

    for node in list(soup.body.children) if soup.body else list(soup.children):
        if getattr(node, 'name', None) is None:
            text = str(node).strip()
            if text:
                p = doc.add_paragraph()
                _clear_paragraph_format(p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _add_inline_html_runs(p, text)
            continue

        if node.name == 'h1':
            p = doc.add_paragraph()
            _clear_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(node.get_text(strip=True))
            _set_run_font(run, size=20, bold=True, color='222222')
        elif node.name == 'h2':
            p = doc.add_paragraph()
            _clear_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(node.get_text(strip=True))
            _set_run_font(run, size=16, bold=True, color='333333')
        elif node.name == 'h3':
            p = doc.add_paragraph()
            _clear_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(node.get_text(strip=True))
            _set_run_font(run, size=13, bold=True, color='444444')
        elif node.name == 'p':
            children = list(node.children)
            if children and all(getattr(c, 'name', None) == 'img' for c in children):
                for child in children:
                    _add_image_to_doc(doc, child.get('src', ''), child.get('alt', ''))
                continue
            has_img = False
            p = doc.add_paragraph()
            _clear_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for child in children:
                if getattr(child, 'name', None) == 'img':
                    has_img = True
                    _add_image_to_doc(doc, child.get('src', ''), child.get('alt', ''))
                    continue
                _add_inline_html_runs(p, child)
            if has_img and not p.text.strip():
                p._element.getparent().remove(p._element)
        elif node.name == 'div' and 'bio-image' in (node.get('class') or []):
            img = node.find('img')
            caption = node.find(class_='bio-image-caption')
            _add_image_to_doc(doc, img.get('src', '') if img else '', caption.get_text(strip=True) if caption else (img.get('alt', '') if img else ''))
        elif node.name == 'blockquote':
            p = doc.add_paragraph()
            _clear_paragraph_format(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.2)
            run = p.add_run(node.get_text(strip=True))
            _set_run_font(run, italic=True, color='555555')
        elif node.name == 'ul':
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                _clear_paragraph_format(p)
                _add_inline_html_runs(p, li)
        elif node.name == 'ol':
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                _clear_paragraph_format(p)
                _add_inline_html_runs(p, li)
        elif node.name == 'hr':
            p = doc.add_paragraph()
            _clear_paragraph_format(p)
            run = p.add_run('_' * 50)
            _set_run_font(run, color='999999')
        else:
            text = node.get_text(' ', strip=True)
            if text:
                p = doc.add_paragraph()
                _clear_paragraph_format(p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _add_inline_html_runs(p, text)

    return doc


def generate_biography_css(markdown_text: str = "", assets: Optional[list] = None) -> str:
    payload = {
        "biography_md": markdown_text,
        "form_data": {},
        "render_target": "web",
        "assets": assets or [],
    }
    skill_path = SKILLS_DIR / "BIO06-layout-css.md"
    try:
        system_prompt = load_skill(str(skill_path))
        result = call_skill("BIO06", system_prompt, payload)
        bio_css = result.get("bio_css") or result.get("css") or ""
        if bio_css:
            return bio_css
    except Exception as exc:
        print(f"[biography] BIO06 generate failed: {exc}")

    return """
.bio-prose-classic {
  max-width: 860px;
  margin: 0 auto;
  padding: 0 6px;
  color: #222;
  font-family: "Noto Serif CJK SC", "Source Han Serif SC", "Songti SC", "SimSun", serif;
  line-height: 1.86;
  font-size: 17px;
  letter-spacing: 0.01em;
  text-align: justify;
}
.bio-prose-classic h1,
.bio-prose-classic h2,
.bio-prose-classic h3 {
  text-align: center;
  margin: 1.4em 0 0.75em;
  line-height: 1.35;
  color: #222;
}
.bio-prose-classic h1 { font-size: 2.1em; letter-spacing: 0.08em; }
.bio-prose-classic h2 { font-size: 1.45em; border-bottom: 1px solid rgba(0,0,0,.08); padding-bottom: .35em; }
.bio-prose-classic h3 { font-size: 1.15em; }
.bio-prose-classic p {
  margin: 0 0 1em 0;
  text-indent: 2em;
  text-align: justify;
  break-inside: avoid;
  overflow-wrap: anywhere;
}
.bio-prose-classic p.no-indent,
.bio-prose-classic .bio-image-caption,
.bio-prose-classic blockquote p,
.bio-prose-classic li p { text-indent: 0; }
.bio-prose-classic blockquote {
  margin: 1em 0;
  padding: .9em 1.1em;
  border-left: 4px solid #d9c09a;
  background: rgba(249, 245, 236, .72);
  color: #555;
}
.bio-prose-classic ul, .bio-prose-classic ol { margin: 0 0 1em 1.4em; }
.bio-prose-classic li { margin: .3em 0; }
.bio-prose-classic .bio-image {
  page-break-inside: avoid;
  break-inside: avoid;
  margin: 1.2em 0;
  max-width: 100%;
}
.bio-prose-classic .bio-image.wrap-left {
  float: left;
  width: min(42%, 340px);
  margin: .25em 1.2em .8em 0;
}
.bio-prose-classic .bio-image img {
  display: block;
  width: auto;
  height: auto;
  max-width: 100%;
  max-height: 50vh;
  object-fit: contain;
  margin: 0 auto;
}
.bio-prose-classic .bio-image-caption {
  text-align: center;
  font-size: .88em;
  color: #666;
  font-style: italic;
  margin-top: .4em;
}
.bio-prose-classic::after { content: ""; display: block; clear: both; }
""".strip()


def convert_images_to_base64(markdown_text: str, sid: str) -> str:
    """将 Markdown 中的图片转换为 base64 嵌入"""
    def replace_with_base64(match):
        alt_text = match.group(1)
        image_path = match.group(2)

        if '/api/memorials/' in image_path or '/assets/' in image_path:
            import re
            asset_match = re.search(r'/assets/([A-Za-z0-9_-]+)', image_path)
            if asset_match:
                asset_id = asset_match.group(1)
                s = session_store.require(sid)
                memorial_id = s.get("form_data", {}).get("memorial_id")
                user_id = s.get("form_data", {}).get("user_id")

                if user_id and memorial_id:
                    asset_meta = next((a for a in core_storage.list_assets(user_id, memorial_id) if a.get("asset_id") == asset_id), None)
                    if asset_meta:
                        stored_name = asset_meta.get("stored_name") or ""
                        asset_path = core_storage.memorial_dir(user_id, memorial_id) / "assets" / stored_name
                        if asset_path.exists():
                            try:
                                with open(asset_path, 'rb') as f:
                                    image_data = f.read()
                                    base64_str = base64.b64encode(image_data).decode('utf-8')
                                    suffix = asset_path.suffix.lower()
                                    mime_map = {
                                        '.jpg': 'image/jpeg',
                                        '.jpeg': 'image/jpeg',
                                        '.png': 'image/png',
                                        '.gif': 'image/gif',
                                        '.webp': 'image/webp'
                                    }
                                    mime_type = mime_map.get(suffix, 'image/jpeg')
                                    return f"![{alt_text}](data:{mime_type};base64,{base64_str})"
                            except Exception as e:
                                print(f"图片转换失败: {asset_path}, error={e}")

        return match.group(0)

    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    return re.sub(pattern, replace_with_base64, markdown_text)


# ── 传记保存、列表、删除 ─────────────────────────────────────────────────


class BioSaveRequest(BaseModel):
    """保存传记请求"""
    sid: str = Field(..., description="session ID")
    title: str = Field(..., description="传记标题（由用户重命名）")
    user_id: str = Field(..., description="用户ID")
    memorial_id: str = Field(..., description="纪念对象ID")


class BioItemResponse(BaseModel):
    """传记项目信息"""
    biography_id: str
    title: str
    deceased_name: str
    created_at: str
    updated_at: str
    word_count: int


@router.post("/save", response_model=dict, summary="保存生成的传记")
def save_biography(req: BioSaveRequest):
    """
    将生成完成的传记保存到memorial存储。
    
    - 从session中提取传记文本和JSON
    - 保存到 data/users/{user_id}/memorials/{memorial_id}/biographies/{biography_id}.md
    - 同时保存元数据到 biographies.json
    """
    try:
        from services import core_storage
        import json
        from datetime import datetime
        
        result = service_manager.get_biography_result(req.sid)
        if result.get("error"):
            raise HTTPException(
                status_code=400,
                detail=result.get("message", "传记还未生成完成"),
            )
        
        # 获取时间戳
        now = datetime.now().isoformat()
        bio_text = result.get("biography_final", "")
        bio_word_count = len(bio_text)
        
        # 生成传记ID
        import uuid
        biography_id = f"bio_{uuid.uuid4().hex[:12]}"
        
        # 保存传记文本
        bio_dir = core_storage.memorial_dir(req.user_id, req.memorial_id) / "biographies"
        import os
        os.makedirs(bio_dir, exist_ok=True)

        bio_file = bio_dir / f"{biography_id}.md"
        with open(bio_file, 'w', encoding='utf-8') as f:
            f.write(bio_text)

        # 保存元数据
        metadata_file = bio_dir / "biographies.json"
        metadata = {}
        if metadata_file.exists():
            with open(metadata_file, 'r', encoding='utf-8') as f:
                try:
                    metadata = json.load(f) or {}
                except Exception:
                    metadata = {}

        s = session_store.require(req.sid)
        deceased_name = s["form_data"].get("deceased_name", "")
        
        metadata[biography_id] = {
            "biography_id": biography_id,
            "title": req.title,
            "deceased_name": deceased_name,
            "created_at": now,
            "updated_at": now,
            "word_count": bio_word_count,
            "quality_assessment": result.get("quality_assessment", {}),
            "info_gaps": result.get("info_gaps", []),
        }
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        print(f"[biography] saved biography_id={biography_id} user={req.user_id} memorial={req.memorial_id}")
        
        return {
            "ok": True,
            "biography_id": biography_id,
            "title": req.title,
            "word_count": bio_word_count,
            "message": f"传记已保存为《{req.title}》",
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"[biography] save error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/list/{user_id}/{memorial_id}", response_model=dict, summary="列表传记历史")
def list_biographies(user_id: str, memorial_id: str):
    """
    获取指定纪念对象的所有生成过的传记列表。
    
    返回按创建时间倒序排列的传记项目。
    """
    try:
        import json
        import os
        
        bio_dir = core_storage.memorial_dir(user_id, memorial_id) / "biographies"
        metadata_file = bio_dir / "biographies.json"
        
        if not metadata_file.exists():
            return {
                "ok": True,
                "biographies": [],
                "count": 0,
            }
        
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f) or {}
        
        # 按created_at倒序排列
        biographies = sorted(
            metadata.values(),
            key=lambda x: x.get("created_at", ""),
            reverse=True
        )
        
        return {
            "ok": True,
            "biographies": biographies,
            "count": len(biographies),
        }
    except Exception as e:
        print(f"[biography] list error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/item/{user_id}/{memorial_id}/{biography_id}", response_model=dict, summary="获取传记内容")
def get_biography_item(user_id: str, memorial_id: str, biography_id: str):
    """
    获取指定传记的完整内容。
    
    返回Markdown格式的传记文本。
    """
    try:
        import os
        
        bio_dir = core_storage.memorial_dir(user_id, memorial_id) / "biographies"
        bio_file = bio_dir / f"{biography_id}.md"
        
        if not bio_file.exists():
            raise HTTPException(status_code=404, detail=f"传记不存在: {biography_id}")
        
        with open(bio_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            "ok": True,
            "biography_id": biography_id,
            "content": content,
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"[biography] get item error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/{user_id}/{memorial_id}/{biography_id}", response_model=dict, summary="删除传记")
def delete_biography(user_id: str, memorial_id: str, biography_id: str):
    """
    删除指定的传记记录。
    
    同时删除文本文件和元数据条目。
    """
    try:
        import json
        import os
        
        bio_dir = core_storage.memorial_dir(user_id, memorial_id) / "biographies"
        bio_file = bio_dir / f"{biography_id}.md"
        metadata_file = bio_dir / "biographies.json"
        
        # 删除文本文件
        if bio_file.exists():
            os.remove(bio_file)
        
        # 更新元数据
        if metadata_file.exists():
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = json.load(f) or {}
            
            if biography_id in metadata:
                del metadata[biography_id]
                
                with open(metadata_file, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        print(f"[biography] deleted biography_id={biography_id} user={user_id} memorial={memorial_id}")
        
        return {
            "ok": True,
            "message": f"传记已删除",
        }
    except Exception as e:
        print(f"[biography] delete error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
